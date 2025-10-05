import os
import tempfile
import streamlit as st
from openai import OpenAI
from io import BytesIO
from docx import Document
from reportlab.pdfgen import canvas

# 🔑 Lê a API Key do ambiente (mais seguro do que hardcode)
API_KEY = os.getenv("OPENAI_API_KEY")

st.set_page_config(page_title="Meeting Summary Generator", page_icon="🎙️", layout="centered")
st.title("🎙️ Meeting Summary Generator")
st.write("Upload an audio file (MP3/WAV) to get a transcription and an AI-generated summary.")

# Campo opcional para testar uma chave sem setar env (só em dev)
with st.expander("🔐 API Key (optional override for local testing)"):
    typed_key = st.text_input("OPENAI_API_KEY", type="password", help="Prefer using the environment variable.")
    if typed_key:
        API_KEY = typed_key

if not API_KEY:
    st.warning("Set the environment variable OPENAI_API_KEY before using the app.")
    st.stop()

# Inicializa cliente
client = OpenAI(api_key=API_KEY)

# ⚙️ Preferências do resumo
col1, col2 = st.columns(2)
with col1:
    idioma = st.selectbox("Resumo em:", ["pt-BR", "en-US"], index=0)
with col2:
    formato = st.selectbox("Formato do resumo:", ["Bullet points", "Texto corrido"], index=0)

# 📂 Upload do áudio
audio_file = st.file_uploader("📂 Upload audio", type=["mp3", "wav"])

# Função para montar o prompt
def build_prompt(texto: str, idioma: str, formato: str) -> str:
    lang = "Português do Brasil" if idioma == "pt-BR" else "English (US)"
    style = "em tópicos (bullet points)" if formato == "Bullet points" else "em texto corrido"
    return f"""
Você é um assistente que cria atas de reunião claras e objetivas.

Tarefa: Resuma a reunião abaixo em {lang}, {style}, com três seções:
1) Resumo executivo (até 3 frases)
2) Decisões tomadas (bullet points)
3) Próximos passos (bullet points)

Seja fiel ao conteúdo. Não invente informações.

### Transcrição
{texto}
"""

# 🚀 Pipeline principal
if audio_file:
    st.audio(audio_file)

    # 1) Transcrição com Whisper
    with st.status("Transcrevendo com Whisper...", expanded=False) as s:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{audio_file.name}") as tmp:
            tmp.write(audio_file.read())
            tmp_path = tmp.name

        try:
            with open(tmp_path, "rb") as f:
                transcription = client.audio.transcriptions.create(
                    model="whisper-1",
                    file=f,
                    response_format="verbose_json"
                )
            texto = transcription.text
            s.update(label="Transcrição concluída ✅", state="complete")
        except Exception as e:
            s.update(label="Falha na transcrição ❌", state="error")
            st.error(f"Erro na transcrição: {e}")
            st.stop()
        finally:
            try:
                os.remove(tmp_path)
            except Exception:
                pass

    # Mostra transcrição
    st.subheader("📝 Transcription")
    st.write(texto if texto.strip() else "_(empty)_")

    # 2) Geração de resumo com GPT
    with st.status("Gerando resumo com GPT...", expanded=False) as s2:
        try:
            prompt = build_prompt(texto, idioma, formato)
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}]
            )
            resumo = resp.choices[0].message.content
            s2.update(label="Resumo criado ✅", state="complete")
        except Exception as e:
            s2.update(label="Falha ao gerar resumo ❌", state="error")
            st.error(f"Erro ao gerar resumo: {e}")
            st.stop()

    # Mostra resumo
    st.subheader("📌 Meeting Summary")
    st.write(resumo)

    # 3) Exportações
    # TXT
    st.download_button(
        "⬇️ Download summary (TXT)",
        data=resumo.encode("utf-8"),
        file_name="meeting_summary.txt",
        mime="text/plain"
    )

    # DOCX
    doc = Document()
    doc.add_heading("Meeting Summary", level=1)
    doc.add_paragraph(resumo)
    buffer_docx = BytesIO()
    doc.save(buffer_docx)
    st.download_button(
        "⬇️ Download summary (DOCX)",
        data=buffer_docx.getvalue(),
        file_name="meeting_summary.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # PDF
    buffer_pdf = BytesIO()
    c = canvas.Canvas(buffer_pdf)
    textobject = c.beginText(40, 800)
    for line in resumo.splitlines():
        textobject.textLine(line)
    c.drawText(textobject)
    c.showPage()
    c.save()
    buffer_pdf.seek(0)
    st.download_button(
        "⬇️ Download summary (PDF)",
        data=buffer_pdf,
        file_name="meeting_summary.pdf",
        mime="application/pdf"
    )
